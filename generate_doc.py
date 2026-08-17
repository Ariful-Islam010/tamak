import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_plantuml_document():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_code_block(doc, code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:left w:val="single" w:sz="18" w:space="0" w:color="0284C7"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(0)
        p_after.paragraph_format.space_after = Pt(8)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("QuitMate - Tobacco Awareness App")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("Complete PlantUML Code Specifications (10 Activity, 10 Sequence, 10 Swimlane Diagrams)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Dictionary of diagrams
    diagrams_data = [
        # Activity Diagrams (7.1 to 7.10)
        ("7. Activity Diagrams", "7.1 Activity Diagram for Manage Profile", """@startuml
title 7.1 Activity Diagram for Manage Profile

start
:Open profile assessment screen;
:Display profile form fields & smoking habit assessment questions;
:Enter personal details (age, gender, institution, smoking per day, stick price);
if (Validate profile data?) then (Valid)
  :Store updated profile details in system;
  :Calculate initial baseline expenditure & stats;
  :Display profile updated confirmation;
else (Invalid / Missing Info)
  :Display field validation error messages;
  :Prompt user to correct input;
  :Correct invalid input fields;
endif
stop
@enduml"""),

        ("7. Activity Diagrams", "7.2 Activity Diagram for Set Quit Date", """@startuml
title 7.2 Activity Diagram for Set Quit Date

start
:Open Quit Plan setup section;
:Prompt target quit date and cigarette/pack cost;
:Select target date and cost per stick/pack;
if (Is selected quit date valid?) then (Valid Future Date)
  :Save target quit date parameter;
  :Calculate daily baseline tobacco cost;
  :Confirm successful quit date setup;
else (Invalid / Past Date)
  :Display invalid date error message;
  :Prompt user to select a valid date;
  :Select valid target date;
endif
stop
@startuml"""),

        ("7. Activity Diagrams", "7.3 Activity Diagram for View Quit Plan", """@startuml
title 7.3 Activity Diagram for View Quit Plan

start
:Navigate to Quit Plan screen;
if (Check server connection) then (Online)
  :Fetch stage milestones & progress calculations from server;
  :Render stage cards, recovery timeline, and personalized advice;
else (Offline)
  :Detect lack of server connectivity;
  :Retrieve locally cached quit plan data;
  :Render quit plan in offline mode;
endif
:View plan stages and health milestones;
stop
@startuml"""),

        ("7. Activity Diagrams", "7.4 Activity Diagram for Daily Check-in", """@startuml
title 7.4 Activity Diagram for Daily Check-in

start
:Initiate daily check-in from dashboard;
:Display craving intensity score (1-10), mood rating, and cigarettes smoked;
:Complete entries and submit check-in;
if (Check existing entry for today) then (New Entry)
  :Validate check-in data;
  :Record check-in record in database;
  :Increment clean streak count;
  :Compute earned Streak Badges;
else (Already Logged Today)
  :Detect check-in already logged today;
  :Prompt update existing check-in entry;
  :Update existing check-in record;
endif
:Display check-in summary and streak animation;
stop
@startuml"""),

        ("7. Activity Diagrams", "7.5 Activity Diagram for Track Saving", """@startuml
title 7.5 Activity Diagram for Track Saving

start
:Open Money Saver screen;
:Calculate accumulated money saved based on quit date & daily cost;
:View total saved amount and existing wishlist items;
:Add new wishlist item with target price;
if (Validate target price?) then (Valid Positive Amount)
  :Save wishlist item to user profile;
  :Compute progress percentage for each wishlist goal;
  :Display active savings progress bar & milestone metrics;
else (Invalid / Zero Amount)
  :Display error message (Invalid target price);
  :Prompt for valid positive amount;
  :Correct target price input;
endif
stop
@startuml"""),

        ("7. Activity Diagrams", "7.6 Activity Diagram for Craving Support", """@startuml
title 7.6 Activity Diagram for Craving Support

start
:Tap red emergency SOS button on dashboard;
:Launch high-priority SOS Support Hub;
:Present craving intervention tools (5-Min Timer, Guided 4-7-8 Breathing, Distraction);
if (Select intervention or exit?) then (Select Tool)
  :Log SOS event for craving analytics;
  :Launch selected emergency intervention mode;
else (Quick Exit)
  :Close SOS Hub;
  :Log short SOS session duration;
  :Return dashboard to normal state;
endif
stop
@startuml"""),

        ("7. Activity Diagrams", "7.7 Activity Diagram for Start Survival Timer", """@startuml
title 7.7 Activity Diagram for Start Survival Timer

start
:Select 5-Minute Survival Timer from SOS hub;
:Initialize 300-second countdown clock with motivational quotes;
repeat
  :Tick down second-by-second;
  :Display rotating survival tips;
backward:Continue countdown;
repeat while (Timer > 0 AND Not Cancelled?) is (Yes)
-> No;

if (Did countdown reach 0:00?) then (Completed)
  :Countdown reaches 0:00;
  :Play completion audio signal;
  :Display success badge dialog;
  :Confirm feeling relieved;
else (Cancelled Early)
  :Tap cancel button before 300s elapse;
  :Stop countdown timer loop;
  :Offer alternative breathing tool;
  :Switch to breathing tool;
endif
stop
@startuml"""),

        ("7. Activity Diagrams", "7.8 Activity Diagram for Breathing Exercise", """@startuml
title 7.8 Activity Diagram for Breathing Exercise

start
:Select Guided Breathing exercise from SOS Hub;
:Check device audio settings;
if (Audio Status?) then (Audio Enabled)
  :Display animated breathing circle with sound cues;
else (Audio Muted)
  :Display visual color-coded animation cues;
endif

repeat
  :Inhale phase (4 seconds - expanding circle);
  :Hold phase (7 seconds - holding circle);
  :Exhale phase (8 seconds - contracting circle);
backward:Next cycle;
repeat while (Target cycles remaining?) is (Yes)
-> No;

:Finish breathing session;
:Play calming audio/visual feedback;
:Experience reduced physical tension;
stop
@startuml"""),

        ("7. Activity Diagrams", "7.9 Activity Diagram for Quick Distraction", """@startuml
title 7.9 Activity Diagram for Quick Distraction

start
:Open Quick Distractions from SOS Hub;
:Present distraction choices (Drink Water, Walk Counter, Sound Distraction);
:Select distraction activity;
if (Switch activity choice?) then (Keep Selected Activity)
  :Execute interactive activity actions;
  :Play sensory sound/visual feedback;
else (Change Activity)
  :Switch activity (e.g. Walking to Sound Distraction);
  :Activate new activity interface;
endif
:Confirm activity completion;
:Record completed distraction task in logs;
stop
@startuml"""),

        ("7. Activity Diagrams", "7.10 Activity Diagram for Access Peer Support Group", """@startuml
title 7.10 Activity Diagram for Access Peer Support Group

start
:Enter Peer Support section;
:Fetch and display recent group chat messages;
:Type a supportive group message;
:Tap Send button;
if (Validate message text?) then (Non-empty Text)
  :Post message payload to backend chat endpoint;
  :Broadcast message to active group members;
  :Render message in real-time chat feed;
else (Empty Text)
  :Detect empty text string;
  :Keep send button disabled;
  :Prompt user to enter valid text;
  :Type valid text message;
endif
stop
@startuml"""),

        # Sequence Diagrams (8.1 to 8.10)
        ("8. Sequence Diagrams", "8.1 Sequence Diagram for Manage Profile", """@startuml
title 8.1 Sequence Diagram for Manage Profile

actor "Tobacco Quitter" as User
participant "App UI" as UI
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Open profile assessment screen
activate UI
UI -> User: 2. Display profile form & assessment questions
User -> UI: 3. Submit profile info (age, gender, stick price, cigarettes/day)
UI -> Backend: 4. Send profile payload for validation
activate Backend

alt Valid Data Flow
    Backend -> DB: 5. Store user profile record
    activate DB
    DB --> Backend: Profile stored successfully
    deactivate DB
    Backend -> Backend: Calculate baseline financial stats
    Backend --> UI: 6. Return success status & baseline stats
    UI --> User: Display profile updated confirmation
else Invalid Data Flow
    Backend --> UI: Return validation errors
    UI --> User: Display error feedback (Prompt correction)
end

deactivate Backend
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.2 Sequence Diagram for Set Quit Date", """@startuml
title 8.2 Sequence Diagram for Set Quit Date

actor "Tobacco Quitter" as User
participant "App UI" as UI
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Navigate to Quit Plan setup screen
activate UI
UI -> User: 2. Prompt for target quit date and stick/pack cost
User -> UI: 3. Select target quit date & cost per stick
UI -> Backend: 4. Send quit date payload for validation
activate Backend

alt Valid Date Selection
    Backend -> Backend: Validate date within allowed range
    Backend -> DB: 5. Save target quit date & daily cost parameters
    activate DB
    DB --> Backend: Data saved confirmation
    deactivate DB
    Backend -> Backend: Compute daily financial baseline
    Backend --> UI: 6. Return quit plan initialization success
    UI --> User: Display setup confirmation message
else Past Date Selected
    Backend --> UI: Return invalid date error
    UI --> User: Prompt to select a valid date
end

deactivate Backend
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.3 Sequence Diagram for View Quit Plan", """@startuml
title 8.3 Sequence Diagram for View Quit Plan

actor "Tobacco Quitter" as User
participant "App UI" as UI
participant "Local Storage (Hive)" as Storage
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Open Quit Plan section
activate UI
UI -> UI: Check network connectivity

alt Online Mode
    UI -> Backend: 2. Fetch quit plan details (Quit Date based)
    activate Backend
    Backend -> DB: Retrieve user milestones & health timeline
    activate DB
    DB --> Backend: Return quit plan data
    deactivate DB
    Backend --> UI: 3. Return quit plan stages & advice
    deactivate Backend
    UI -> Storage: Cache updated quit plan locally
else Offline Mode
    UI -> Storage: 2b. Retrieve locally cached quit plan
    activate Storage
    Storage --> UI: 2c. Return cached quit plan payload
    deactivate Storage
end

UI --> User: 4. Display stage cards, health recovery timeline, and guidance
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.4 Sequence Diagram for Daily Check-in", """@startuml
title 8.4 Sequence Diagram for Daily Check-in

actor "Tobacco Quitter" as User
participant "App UI" as UI
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Open Daily Check-in screen
activate UI
UI -> User: 2. Prompt for craving score (1-10), mood, and tobacco count
User -> UI: 3. Enter check-in details and submit
UI -> Backend: 4. Send check-in payload
activate Backend

alt New Check-in Today
    Backend -> DB: 5. Store daily check-in entry
    activate DB
    DB --> Backend: Check-in saved
    deactivate DB
    Backend -> Backend: Update clean streak & compute Streak Badges
    Backend --> UI: 6. Return check-in summary & streak update
    UI --> User: Display check-in summary & streak animation
else Duplicate Check-in Today
    Backend -> DB: 4b. Update existing check-in record for today
    activate DB
    DB --> Backend: Check-in updated
    deactivate DB
    Backend --> UI: Return updated entry summary
    UI --> User: Display updated check-in summary
end

deactivate Backend
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.5 Sequence Diagram for Track Saving", """@startuml
title 8.5 Sequence Diagram for Track Saving

actor "Tobacco Quitter" as User
participant "App UI" as UI
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Navigate to Money Saver screen
activate UI
UI -> Backend: 2. Fetch user financial baseline & quit date
activate Backend
Backend -> DB: Retrieve quit date & daily stick cost
activate DB
DB --> Backend: Financial baseline data
deactivate DB
Backend -> Backend: Calculate accumulated money saved
Backend --> UI: 3. Display total money saved & active wishlist items
deactivate Backend

User -> UI: 4. Add new wishlist item (Title, Target Price)

alt Valid Target Price
    UI -> Backend: 5. Send wishlist goal payload
    activate Backend
    Backend -> DB: Store new wishlist goal record
    activate DB
    DB --> Backend: Goal saved confirmation
    deactivate DB
    Backend -> Backend: Compute completion percentage
    Backend --> UI: 6. Return updated wishlist & progress %
    UI --> User: Render active savings progress bar
else Invalid Target Price
    UI --> User: Display error (Invalid target amount)
end

deactivate Backend
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.6 Sequence Diagram for Craving Support", """@startuml
title 8.6 Sequence Diagram for Craving Support

actor "Tobacco Quitter" as User
participant "Dashboard UI" as UI
participant "SOS Support Hub" as SOS
participant "Backend System" as Backend
database "Database" as DB

User -> UI: 1. Tap emergency SOS button
activate UI
UI -> SOS: 2. Launch high-priority SOS Support Hub
activate SOS
SOS -> User: 3. Present options (Survival Timer, 4-7-8 Breathing, Distractions)

alt Select Intervention Tool
    User -> SOS: 5. Select emergency tool option
    SOS -> Backend: 4. Log SOS event & selected mode
    activate Backend
    Backend -> DB: Store SOS log record
    activate DB
    DB --> Backend: Log saved
    deactivate DB
    Backend --> SOS: Ack SOS log
    deactivate Backend
    SOS --> User: Launch selected craving intervention module
else Quick Exit Flow
    User -> SOS: 5a. Close SOS hub
    SOS -> Backend: Log brief SOS event duration
    SOS --> UI: 5c. Return to normal dashboard view
end

deactivate SOS
deactivate UI
@startuml"""),

        ("8. Sequence Diagrams", "8.7 Sequence Diagram for Start Survival Timer", """@startuml
title 8.7 Sequence Diagram for Start Survival Timer

actor "Tobacco Quitter" as User
participant "SOS Hub UI" as SOS
participant "Survival Timer UI" as Timer
participant "Audio / Haptic Engine" as Media

User -> SOS: 1. Select 5-Minute Survival Timer
activate SOS
SOS -> Timer: 2. Initialize 300-second countdown clock
activate Timer
Timer -> User: Display timer (300s) & motivational quotes

alt Countdown Completion Flow
    loop Every Second (for 300 seconds)
        Timer -> Timer: 3. Tick down 1 sec & update survival tips
    end
    Timer -> Media: 5. Trigger completion sound signal
    activate Media
    Media --> Timer: Audio playback finished
    deactivate Media
    Timer -> User: Display success completion badge & prompt feedback
    User -> Timer: 6. Confirm feeling relieved
else Early Cancellation Flow
    User -> Timer: 3a. Tap cancel timer button
    Timer -> Timer: 3b. Stop countdown clock
    Timer -> User: Offer guided breathing tool alternative
    User -> Timer: 3c. Switch tool mode
end

deactivate Timer
deactivate SOS
@startuml"""),

        ("8. Sequence Diagrams", "8.8 Sequence Diagram for Breathing Exercise", """@startuml
title 8.8 Sequence Diagram for Breathing Exercise

actor "Tobacco Quitter" as User
participant "SOS Hub UI" as SOS
participant "Breathing UI" as UI
participant "Audio Engine" as Audio

User -> SOS: 1. Select Guided Breathing exercise
activate SOS
SOS -> UI: 2. Launch 4-7-8 Breathing interface
activate UI

alt Audio Enabled
    UI -> Audio: Start calming background rhythm
    activate Audio
    Audio --> UI: Audio playing
    deactivate Audio
else Audio Muted
    UI -> UI: 2b. Enable high-contrast visual cues mode
end

loop Designated Breathing Cycles (e.g. 4 Cycles)
    UI -> User: 3. Phase 1: Inhale (4 seconds) - Expand circle
    UI -> User: Phase 2: Hold breath (7 seconds) - Hold circle
    UI -> User: Phase 3: Exhale (8 seconds) - Contract circle
end

UI -> Audio: 5. Play session completion chime
UI --> User: Display session completed summary & calmness rating
deactivate UI
deactivate SOS
@startuml"""),

        ("8. Sequence Diagrams", "8.9 Sequence Diagram for Quick Distraction", """@startuml
title 8.9 Sequence Diagram for Quick Distraction

actor "Tobacco Quitter" as User
participant "SOS Hub UI" as SOS
participant "Distraction UI" as UI
participant "Backend System" as Backend
database "Database" as DB

User -> SOS: 1. Open Quick Distractions
activate SOS
SOS -> UI: 2. Present options (Water Tracker, Walk Counter, Sound Distraction)
activate UI
User -> UI: 3. Select an activity (e.g., Drink Water)

alt Standard Completion Flow
    UI -> User: 4. Execute activity action & play interactive sound
    User -> UI: 5. Confirm activity completion
    UI -> Backend: 6. Record completed distraction activity log
    activate Backend
    Backend -> DB: Save distraction log record
    activate DB
    DB --> Backend: Log stored
    deactivate DB
    Backend --> UI: Return completion confirmation
    deactivate Backend
    UI --> User: Display positive feedback badge
else Change Activity Flow
    User -> UI: 3a. Switch activity selection (e.g. Walk to Sound Player)
    UI -> UI: 3b. Activate sound player interface
    UI -> User: 3c. Proceed with sound distraction activity
end

deactivate UI
deactivate SOS
@startuml"""),

        ("8. Sequence Diagrams", "8.10 Sequence Diagram for Access Peer Support Group", """@startuml
title 8.10 Sequence Diagram for Access Peer Support Group

actor "Quitter / Peer Member" as User
participant "Chat UI" as UI
participant "Chat Server / API" as Server
database "Database" as DB

User -> UI: 1. Open Peer Support Group Chat
activate UI
UI -> Server: 2. Fetch recent group chat history
activate Server
Server -> DB: Retrieve chat message logs
activate DB
DB --> Server: Return message records
deactivate DB
Server --> UI: Return recent messages array
deactivate Server
UI --> User: Display live chatroom feed

User -> UI: 3. Type supportive message & tap Send

alt Non-empty Message Flow
    UI -> Server: 4. Post message payload (sender_id, content)
    activate Server
    Server -> Server: Validate content (filter inappropriate words)
    Server -> DB: 5. Store message record in database
    activate DB
    DB --> Server: Message saved
    deactivate DB
    Server --> UI: 6. Broadcast message to chatroom clients
    deactivate Server
    UI --> User: Render message in chat feed real-time
else Empty Message Flow
    UI -> UI: 4a. Detect empty message text
    UI --> User: 4b. Keep send button disabled / prompt text entry
end

deactivate UI
@startuml"""),

        # Swimlane Diagrams (9.1 to 9.10)
        ("9. Swimlane Diagrams", "9.1 Swimlane Diagram for Manage Profile", """@startuml
title 9.1 Swimlane Diagram for Manage Profile

|Tobacco Quitter|
start
:Open profile assessment screen;

|App Interface|
:Render profile fields and smoking habit questions;

|Tobacco Quitter|
:Enter personal details, cigarettes/day, stick price;
:Submit profile assessment form;

|Backend System|
if (Validate input details?) then (Valid)
  |Database|
  :Save updated profile record;
  |Backend System|
  :Compute baseline smoking expenditure;
  |App Interface|
  :Show profile updated confirmation message;
else (Invalid)
  |App Interface|
  :Show input validation error messages;
  |Tobacco Quitter|
  :Correct invalid fields and resubmit;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.2 Swimlane Diagram for Set Quit Date", """@startuml
title 9.2 Swimlane Diagram for Set Quit Date

|Tobacco Quitter|
start
:Open Quit Plan setup section;

|App Interface|
:Prompt target quit date and cigarette cost;

|Tobacco Quitter|
:Select target date and cost per stick/pack;
:Submit quit date form;

|Backend System|
if (Validate target date?) then (Valid Date)
  |Database|
  :Save target quit date & cost baseline;
  |Backend System|
  :Calculate daily financial baseline cost;
  |App Interface|
  :Display successful setup confirmation;
else (Past Date)
  |App Interface|
  :Display error message (Past date selected);
  |Tobacco Quitter|
  :Select valid target date and resubmit;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.3 Swimlane Diagram for View Quit Plan", """@startuml
title 9.3 Swimlane Diagram for View Quit Plan

|Tobacco Quitter|
start
:Navigate to Quit Plan screen;

|App Interface|
if (Network Connectivity?) then (Online)
  |Backend System|
  :Fetch stage milestones and recovery progress;
  |Database|
  :Retrieve user quit plan data;
  |Backend System|
  :Format stage cards and advice payload;
  |App Interface|
  :Render stage cards, timeline, and advice;
else (Offline)
  |Local Cache (Hive)|
  :Retrieve locally cached quit plan;
  |App Interface|
  :Display quit plan in offline mode;
endif

|Tobacco Quitter|
:View stage milestones and progress;
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.4 Swimlane Diagram for Daily Check-in", """@startuml
title 9.4 Swimlane Diagram for Daily Check-in

|Tobacco Quitter|
start
:Initiate daily check-in from dashboard;

|App Interface|
:Prompt craving intensity (1-10), mood, cigarettes count;

|Tobacco Quitter|
:Enter check-in details and tap submit;

|Backend System|
if (Check-in status for today?) then (New Entry)
  |Database|
  :Record daily check-in entry;
  |Backend System|
  :Increment clean streak & compute Streak Badges;
  |App Interface|
  :Show check-in summary & streak animation;
else (Duplicate Entry)
  |App Interface|
  :Allow updating existing check-in entry;
  |Backend System|
  :Update check-in record in database;
  |App Interface|
  :Show updated check-in summary;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.5 Swimlane Diagram for Track Saving", """@startuml
title 9.5 Swimlane Diagram for Track Saving

|Tobacco Quitter|
start
:Navigate to Money Saver screen;

|App Interface|
:Request financial calculations;

|Backend System|
:Calculate total saved money based on quit date & cost;

|App Interface|
:Display total saved amount and wishlist list;

|Tobacco Quitter|
:Enter new wishlist item and target price;

|Backend System|
if (Validate target price?) then (Valid Price)
  |Database|
  :Store new wishlist item record;
  |Backend System|
  :Compute progress percentage for goals;
  |App Interface|
  :Display updated progress bar & goal percentage;
else (Invalid Price)
  |App Interface|
  :Show target price error message;
  |Tobacco Quitter|
  :Correct price input and resubmit;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.6 Swimlane Diagram for Craving Support", """@startuml
title 9.6 Swimlane Diagram for Craving Support

|Tobacco Quitter|
start
:Tap emergency SOS button on dashboard;

|App Interface|
:Launch high-priority SOS Support Hub;
:Present intervention options (Timer, Breathing, Distraction);

|Tobacco Quitter|
if (Choose action?) then (Select Tool)
  |Backend System|
  :Log SOS event and selected tool mode;
  |Database|
  :Store SOS analytics log record;
  |App Interface|
  :Launch selected intervention screen;
else (Quick Exit)
  |App Interface|
  :Dismiss SOS Hub;
  |Backend System|
  :Log short SOS interaction duration;
  |App Interface|
  :Return dashboard to normal state;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.7 Swimlane Diagram for Start Survival Timer", """@startuml
title 9.7 Swimlane Diagram for Start Survival Timer

|Tobacco Quitter|
start
:Select 5-Minute Survival Timer from SOS Hub;

|App Interface|
:Initialize 300-second countdown timer with quotes;

|Tobacco Quitter|
while (Countdown active & not cancelled?) is (Running)
  |App Interface|
  :Tick down 1 second & display survival tips;
  |Tobacco Quitter|
endwhile (Finished or Cancelled)

|App Interface|
if (Reason for stopping?) then (Timer Reached 0:00)
  :Play completion audio signal;
  :Show success badge dialog;
  |Tobacco Quitter|
  :Confirm craving urge subsided;
else (User Cancelled Early)
  |App Interface|
  :Stop countdown clock;
  :Offer alternative guided breathing tool;
  |Tobacco Quitter|
  :Switch to breathing tool;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.8 Swimlane Diagram for Breathing Exercise", """@startuml
title 9.8 Swimlane Diagram for Breathing Exercise

|Tobacco Quitter|
start
:Select Guided Breathing from SOS Hub;

|App Interface|
if (Device Audio Setting?) then (Enabled)
  :Initialize animated breathing circle with sound cues;
else (Muted)
  :Initialize high-contrast visual animation cues;
endif

|Tobacco Quitter|
repeat
  |App Interface|
  :Guide Inhale phase (4 seconds - Expanding circle);
  :Guide Hold phase (7 seconds - Holding circle);
  :Guide Exhale phase (8 seconds - Contracting circle);
  |Tobacco Quitter|
backward:Repeat cycle;
repeat while (Session duration active?) is (Yes)
-> No;

|App Interface|
:Finish breathing session & play calming feedback;

|Tobacco Quitter|
:Confirm reduced craving anxiety & tension;
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.9 Swimlane Diagram for Quick Distraction", """@startuml
title 9.9 Swimlane Diagram for Quick Distraction

|Tobacco Quitter|
start
:Open Quick Distractions menu from SOS Hub;

|App Interface|
:Display distraction choices (Water Tracker, Walk Counter, Sound player);

|Tobacco Quitter|
:Select desired distraction activity;

|App Interface|
if (Activity selection state?) then (Execute Activity)
  :Execute activity actions & play sensory feedback;
  |Tobacco Quitter|
  :Perform sensory activity & confirm completion;
  |Backend System|
  :Record completed distraction task log;
  |Database|
  :Store distraction task log;
  |App Interface|
  :Show task completed feedback;
else (Switch Activity)
  |App Interface|
  :Switch to new activity view (e.g. Sound Player);
  |Tobacco Quitter|
  :Proceed with newly selected distraction;
endif

|Tobacco Quitter|
stop
@startuml"""),

        ("9. Swimlane Diagrams", "9.10 Swimlane Diagram for Access Peer Support Group", """@startuml
title 9.10 Swimlane Diagram for Access Peer Support Group

|Tobacco Quitter / Peer Member|
start
:Open Peer Support Group Chat screen;

|App Interface|
:Request recent message history;

|Backend Chat Server|
:Retrieve chat messages from database;
|Database|
:Fetch message records;
|App Interface|
:Render message history in chat feed;

|Tobacco Quitter / Peer Member|
:Type supportive message and tap send;

|App Interface|
if (Validate message content?) then (Valid Text)
  |Backend Chat Server|
  :Validate & sanitize message payload;
  |Database|
  :Store message record;
  |Backend Chat Server|
  :Broadcast message payload to connected clients;
  |App Interface|
  :Append new message in real-time chat feed;
else (Empty Text)
  |App Interface|
  :Disable send button / show validation error;
  |Tobacco Quitter / Peer Member|
  :Enter valid text message and resubmit;
endif

|Tobacco Quitter / Peer Member|
stop
@startuml""")
    ]

    current_section = ""
    for sec, title, code in diagrams_data:
        if sec != current_section:
            current_section = sec
            # Add Section Heading (Heading 1)
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(18)
            p_sec.paragraph_format.space_after = Pt(8)
            p_sec.paragraph_format.keep_with_next = True
            r_sec = p_sec.add_run(sec)
            r_sec.font.name = 'Calibri'
            r_sec.font.size = Pt(18)
            r_sec.font.bold = True
            r_sec.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep navy blue

        # Add Diagram Title (Heading 2)
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(12)
        p_t.paragraph_format.space_after = Pt(4)
        p_t.paragraph_format.keep_with_next = True
        r_t = p_t.add_run(title)
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(14)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x02, 0x84, 0xC7) # Sky blue accent

        # Add PlantUML Code Block
        add_code_block(doc, code)

    # Save outputs
    output_path1 = r"c:\Users\hp\OneDrive\Desktop\App\QuitMate_PlantUML_Diagrams.docx"
    output_path2 = r"c:\Users\hp\OneDrive\Desktop\QuitMate_PlantUML_Diagrams.docx"
    
    doc.save(output_path1)
    try:
        doc.save(output_path2)
        print(f"Successfully saved to {output_path2}")
    except Exception as e:
        print(f"Could not save to desktop directly: {e}")
        
    print(f"Successfully created docx document at: {output_path1}")

if __name__ == "__main__":
    create_plantuml_document()
